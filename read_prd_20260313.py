from docx import Document

doc = Document(r'D:\RMS_Siprahub\docs\RMS_PRD_20260313.docx')

with open(r'd:\RMS_Siprahub\prd_dump_20260313.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(f"[{para.style.name}] {para.text}\n")
    
    for table in doc.tables:
        f.write("\n--- TABLE ---\n")
        for row in table.rows:
            f.write(' | '.join([cell.text for cell in row.cells]) + '\n')
        f.write("--- END TABLE ---\n")

print("Done. Output in prd_dump_20260313.txt")
