"""
Complete RMS PRD Generator - Production Quality
Includes:
- Comprehensive content for all 8 sections
- Text-based Table of Contents (for manual conversion)
- Proper table formatting (10%, 25%, 65%)
- Phase 1 and Phase 2 scope clearly defined
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level)
    run = heading.runs[0]
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 71, 136)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(46, 92, 138)
    elif level == 3:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(70, 130, 180)
    
    run.font.bold = True
    return heading

def set_column_width(table, col_idx, width_pct):
    """Set column width as percentage"""
    for row in table.rows:
        cell = row.cells[col_idx]
        cell.width = Inches(6.5 * width_pct / 100)

def create_complete_prd():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(13)
    
    # ============== TITLE PAGE ==============
    title = doc.add_paragraph('Product Requirements Document')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph('Resource Management System (RMS)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    for label, value in [
        ('Document Version:', '1.0'),
        ('Date:', 'February 16, 2026'),
        ('Project Timeline:', '2 Weeks (Feb 16 - March 2, 2026)'),
        ('Confidentiality:', 'Internal Use Only')
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'{label} ').bold = True
        p.add_run(value)
    
    doc.add_page_break()
    
    # ============== TABLE OF CONTENTS (TEXT) ==============
    toc_heading = doc.add_paragraph('TABLE OF CONTENTS')
    toc_heading.runs[0].font.size = Pt(18)
    toc_heading.runs[0].font.bold = True
    toc_heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph()
    
    toc_items = [
        '1. Executive Summary',
        '   1.1 Project Overview',
        '   1.2 Problem Statement',
        '   1.3 Business Objectives',
        '',
        '2. User Personas',
        '   2.1 Back Office Recruiter',
        '   2.2 Admin (Profile Review & Client Submission)',
        '   2.3 Client Portal User (Phase 2)',
        '',
        '3. User Stories & Acceptance Criteria',
        '   US-001: Resource Request Creation',
        '   US-002: Candidate Addition (21-Field Form)',
        '   US-003: Admin Profile Review',
        '   US-004: Client Submission',
        '   US-005: Onboarding Workflow',
        '   US-006: Exit Processing',
        '   US-007: Dashboard Visibility',
        '   US-008: SOW Tracker',
        '   US-009: Backfill Automation',
        '   US-010: Communication Logging',
        '',
        '4. Functional Requirements',
        '   4.1 Authentication & Role-Based Access Control',
        '   4.2 Dashboard & Real-Time Metrics',
        '   4.3 Job Profile Management',
        '   4.4 Resource Request Workflow',
        '   4.5 Recruiter Pipeline (21-Field Form)',
        '   4.6 Admin Review & Client Submission',
        '   4.7 Onboarding Workflow',
        '   4.8 Exit Management & Backfill',
        '   4.9 SOW Tracker (Manual Entry)',
        '',
        '5. Candidate Fields Specification',
        '   5.1 21-Field Breakdown',
        '   5.2 Field Validation Rules',
        '',
        '6. Scope & Boundaries',
        '   6.1 Phase 1: Core Platform (2 Weeks)',
        '   6.2 Phase 2: Advanced Features (Future)',
        '   6.3 Out of Scope',
        '',
        '7. Assumptions & Dependencies',
        '   7.1 Technical Assumptions',
        '   7.2 Business Assumptions',
        '   7.3 External Dependencies',
        '',
        '8. Success Criteria & KPIs',
        '   8.1 Quantitative Metrics',
        '   8.2 Qualitative Metrics',
        '   8.3 User Adoption Targets',
    ]
    
    for item in toc_items:
        if item:
            doc.add_paragraph(item)
        else:
            doc.add_paragraph()
    
    p_note = doc.add_paragraph()
    p_note.add_run('Note: ').bold = True
    p_note.add_run('To convert to a table, select all TOC items, then Insert → Table → Convert Text to Table (2 columns).')
    p_note.runs[1].font.italic = True
    
    doc.add_page_break()
    
    # ============== 1. EXECUTIVE SUMMARY ==============
    add_heading(doc, '1. Executive Summary', 1)
    
    add_heading(doc, '1.1 Project Overview', 2)
    doc.add_paragraph(
        'The Resource Management System (RMS) is a comprehensive web-based platform designed to replace the current '
        'manual Excel-based workflow for managing staff augmentation operations. The system will centralize tracking '
        'of resource requests, candidate sourcing, admin review, client submission, onboarding, and lifecycle '
        'management from initial request through exit.'
    )
    
    add_heading(doc, '1.2 Problem Statement', 2)
    doc.add_paragraph(
        'The organization currently manages approximately 300+ active resources using multiple Excel spreadsheets '
        '(Resource Data, SOW Tracker, Category Data, Recruiter Pipeline). This manual process results in:'
    )
    
    for issue in [
        'Data inconsistencies due to manual entry (e.g., role name variations like "Java Developer" vs "Java Dev")',
        'No centralized profile storage - requires email searches to retrieve candidate resumes',
        'Client visibility into high attrition rates when sharing Excel files directly',
        'Difficult billing reconciliation with client systems due to Jira username mapping issues',
        'No audit trail for request origins (email vs chat) or client communications',
        'Inability to scale beyond 300-500 resources without exponential manual effort',
        'Backfill requests created manually with no linkage to original exits/rejections'
    ]:
        doc.add_paragraph(issue, style='List Bullet')
    
    add_heading(doc, '1.3 Business Objectives', 2)
    doc.add_paragraph('The RMS aims to:')
    
    for obj in [
        'Centralize all staff augmentation operations in one secure, web-based platform',
        'Reduce data entry errors by 90% through structured forms and validation rules',
        'Provide real-time visibility into resource pipeline, onboarded headcount, and attrition trends',
        'Streamline recruiter-to-admin handoff for client profile submissions',
        'Automate backfill request creation for exits and client rejections',
        'Enable scalability to support500+ concurrent resources without performance degradation',
        'Create comprehensive audit trail for all request sources and client communications'
    ]:
        doc.add_paragraph(obj, style='List Bullet')
    
    # ============== 2. USER PERSONAS ==============
    add_heading(doc, '2. User Personas', 1)
    
    add_heading(doc, '2.1 Back Office Recruiter', 2)
    p = doc.add_paragraph()
    p.add_run('Role: ').bold = True
    p.add_run('Handles candidate sourcing, screening, and initial submission to admin')
    
    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True
    for resp in [
        'Create resource requests when client requirements are received',
        'Source candidates from vendors (WRS, GFM) or internal database',
        'Add candidates with complete 21-field profile form',
        'Upload candidate resumes (PDF/DOCX)',
        'Update candidate status through pipeline (Interview Scheduled, Submitted to Admin, etc.)',
        'Mark requests as "With Admin" when ready for profile review'
    ]:
        doc.add_paragraph(resp, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Pain Points (Current):').bold = True
    for pain in [
        'Manually copying data across multiple Excel sheets',
        'No dropdown validations - frequent typos in role names and statuses',
        'Cannot track which recruiter owns which candidate',
        'Searching emails to find candidate resumes submitted weeks ago'
    ]:
        doc.add_paragraph(pain, style='List Bullet')
    
    add_heading(doc, '2.2 Admin (Profile Review & Client Submission)', 2)
    p = doc.add_paragraph()
    p.add_run('Role: ').bold = True
    p.add_run('Reviews recruiter-submitted profiles and sends them to clients')
    
    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True
    for resp in [
        'View all requests marked "With Admin"',
        'Validate candidate profiles (email format, JD alignment, resume quality)',
        'Reject profiles back to recruiter with feedback if needed',
        'Generate professional email templates for client submission',
        'Log communication details (date sent, client contact)',
        'Update status to "With Client" after submission',
        'Process client selections and rejections'
    ]:
        doc.add_paragraph(resp, style='List Bullet')
    
    add_heading(doc, '2.3 Client Portal User (Phase 2)', 2)
    doc.add_paragraph('Future persona - not part of Phase 1 scope. See Section 6.2 for Phase 2 features.')
    
    # ============== 3. USER STORIES ==============
    add_heading(doc, '3. User Stories & Acceptance Criteria', 1)
    
    stories = [
        ('US-001', 'Resource Request Creation',
         'As a Back Office Recruiter, I want to create resource requests with auto-generated IDs so that I can track client requirements systematically.',
         ['Request ID is auto-generated (format: REQ-YYYYMMDD-XXX)',
          'Job profile is selected from dropdown',
          'Request source (Email/Chat) is captured',
          'Priority (Urgent/High/Medium/Low) can be set',
          'Multi-position requests are supported',
          'Request is saved and visible in list view']),
        
        ('US-002', 'Candidate Addition (21-Field Form)',
         'As a Back Office Recruiter, I want to add candidates with complete profiles so that Admin can review all necessary information.',
         ['All 21 fields can be filled (see Section 5)',
          'Resume can be uploaded (PDF/DOCX, max 5MB)',
          'Owner (recruiter) is assigned from dropdown',
          'Candidate status defaults to "New"',
          'Candidate is linked to request ID',
          'Form validates required fields before save']),
        
        ('US-003', 'Admin Profile Review',
         'As an Admin, I want to review recruiter-submitted profiles so that I only send quality candidates to clients.',
         ['Can view all requests "With Admin"',
          'Can see complete candidate profile with resume',
          'Can reject to recruiter with reason',
          'Can approve and proceed to client submission',
          'Validation checklist is visible (email, JD match, resume quality)']),
        
        ('US-004', 'Client Submission',
         'As an Admin, I want to generate professional emails for client submission so that the process is standardized.',
         ['Email template is auto-generated with candidate details',
          'Can download candidate profile as PDF',
          'Can log communication date and client contact',
          'Status updates to "With Client" after submission',
          'Audit trail is maintained']),
        
        ('US-005', 'Onboarding Workflow',
         'As an Admin, I want to capture onboarding details so that billing and client systems are synchronized.',
         ['Can mark candidate as onboarded',
          'Billing start date is captured',
          'Client email ID is captured',
          'Client Jira username is captured',
          'Status updates to "Onboarded"',
          'Dashboard metrics update in real-time']),
        
        ('US-006', 'Exit Processing',
         'As an Admin, I want to process exits with reasons so that attrition trends can be analyzed.',
         ['Can select exit reason from dropdown (6 options)',
          'Last working day is captured',
          'Exit notes can be added',
          'Replacement required flag can be set',
          'Backfill request is auto-created if replacement needed',
          'Status updates to "Exit"']),
        
        ('US-007', 'Dashboard Visibility',
         'As a user, I want to see real-time metrics so that I can track pipeline health.',
         ['Total requests count is displayed',
          'Onboarded count is displayed',
          'Awaiting onboarding count (status = "With Client") is shown',
          'To be shared count (status = "With Admin") is shown',
          'Role-wise breakdown chart is displayed',
          'Technology distribution is visible']),
        
        ('US-008', 'SOW Tracker',
         'As an Admin, I want to track SOW details so that headcount can be reconciled with contracts.',
         ['Can manually enter SOW details',
          'Can link request IDs to SOW',
          'Can view SOW list in table format',
          'Can see request count per SOW']),
        
        ('US-009', 'Backfill Automation',
         'As an Admin, I want backfill requests to be auto-created so that replacements are not forgotten.',
         ['Backfill is created when exit has "replacement required" = Yes',
          'Backfill is created when client rejection has "replacement required" = Yes',
          'Backfill request inherits job profile from original request',
          'Backfill is linked to original request ID']),
        
        ('US-010', 'Communication Logging',
         'As an Admin, I want to log all client communications so that there is an audit trail.',
         ['Can log communication date',
          'Can log client contact person',
          'Can log communication type (email/call/meeting)',
          'All logs are timestamped and linked to request']),
    ]
    
    for story_id, title, description, criteria in stories:
        p = doc.add_paragraph()
        p.add_run(f'{story_id}: {title}').bold = True
        
        doc.add_paragraph(description)
        
        p = doc.add_paragraph()
        p.add_run('Acceptance Criteria:').bold = True
        for criterion in criteria:
            doc.add_paragraph(criterion, style='List Bullet')
        
        doc.add_paragraph()
    
    # ============== 4. FUNCTIONAL REQUIREMENTS ==============
    add_heading(doc, '4. Functional Requirements', 1)
    
    modules = [
        ('4.1 Authentication & Role-Based Access Control',
         ['User login with email and password',
          'Two roles: Back Office (Recruiter) and Admin',
          'Permission-based feature access (e.g., only Admin can "Send to Client")',
          'Session persistence across page refresh',
          'Logout functionality']),
        
        ('4.2 Dashboard & Real-Time Metrics',
         ['Total requests metric',
          'Onboarded count',
          'Awaiting onboarding count (status = "With Client")',
          'To be shared count (status = "With Admin")',
          'Role-wise breakdown (bar/pie chart)',
          'Technology distribution chart',
          'Status filter dropdown',
          'Attrition rate trend graph (Phase 1 basic, Phase 2 advanced)',
          'Average time to onboard metric',
          'Date range filter']),
        
        ('4.3 Job Profile Management',
         ['Create job profile (role name, technology, experience)',
          'Edit job profile',
          'Delete job profile (with validation - cannot delete if linked to requests)',
          'List job profiles (paginated)',
          'Duplicate validation (same role name cannot exist twice)']),
        
        ('4.4 Resource Request Workflow',
         ['Create request with auto-generated ID (REQ-YYYYMMDD-XXX)',
          'Select job profile from dropdown',
          'Capture request source (Email/Chat)',
          'Set priority (Urgent/High/Medium/Low)',
          'Support multi-position requests',
          'View request list (paginated)',
          'Filter requests by status, priority, job profile',
          'Search by request ID or candidate name']),
        
        ('4.5 Recruiter Pipeline (21-Field Form)',
         ['Add candidate with 21 fields (see Section 5)',
          'Owner assignment (recruiter dropdown)',
          'Vendor field (WRS/GFM/Internal)',
          'Interview date & time picker',
          'Candidate status dropdown (9 options)',
          'CTC fields (current/expected)',
          'Location fields (current/work location)',
          'Notice period field',
          'Remarks (multi-line text)',
          'Upload resume (PDF/DOCX, max 5MB)',
          'Edit candidate details (all 21 fields)',
          'View candidate list for request',
          'Update request status to "With Admin"']),
        
        ('4.6 Admin Review & Client Submission',
         ['View requests "With Admin"',
          'Validation checklist (email format, JD alignment, resume quality)',
          'Download candidate profile as PDF',
          'Reject to Back Office with reason',
          'Generate email template for client',
          'Update status to "With Client"',
          'Log communication details (date, client contact, type)']),
        
        ('4.7 Onboarding Workflow',
         ['Mark as onboarded',
          'Capture billing start date',
          'Capture client email ID',
          'Capture client Jira username',
          'Update status to "Onboarded"',
          'Client rejection flow (with reason)',
          'Replacement required choice',
          'Auto-create backfill on rejection if replacement needed']),
        
        ('4.8 Exit Management & Backfill',
         ['Process exit (capture reason & last working day)',
          'Exit reason dropdown (6 options: Better Offer, Personal, Performance, Client End, Project End, Other)',
          'Last working day date picker',
          'Exit notes (multi-line text)',
          'Replacement required choice',
          'Auto-create backfill on exit if replacement needed',
          'Update status to "Exit"']),
        
        ('4.9 SOW Tracker (Manual Entry)',
         ['Manual SOW entry form',
          'Link SOW to request IDs',
          'View SOW list (table)',
          'Display request count per SOW']),
    ]
    
    for heading, features in modules:
        add_heading(doc, heading, 2)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
    
    # ============== 5. CANDIDATE FIELDS SPECIFICATION ==============
    add_heading(doc, '5. Candidate Fields Specification', 1)
    add_heading(doc, '5.1 21-Field Breakdown', 2)
    
    doc.add_paragraph().add_run('Complete field specification for candidate form:').bold = True
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Field Name'
    hdr_cells[2].text = 'Description / Validation Rules'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    fields = [
        ('1', 'Owner', 'Recruiter assigned (dropdown from users table)'),
        ('2', 'Date', 'Date candidate added (auto-populated, format: DD-MMM-YYYY)'),
        ('3', 'Vendor', 'Sourcing vendor (dropdown: WRS, GFM, Internal)'),
        ('4', 'Interview Date', 'Scheduled interview date (date picker)'),
        ('5', 'Candidate Name', 'Full name (text, max 100 chars, required)'),
        ('6', 'Email', 'Email address (email validation, required)'),
        ('7', 'Phone', 'Contact number (10 digits, numeric)'),
        ('8', 'Current Company', 'Current employer name (text, max 100 chars)'),
        ('9', 'Current CTC', 'Current annual salary (numeric, in lakhs)'),
        ('10', 'Expected CTC', 'Expected annual salary (numeric, in lakhs)'),
        ('11', 'Current Location', 'Current city/location (text, max 50 chars)'),
        ('12', 'Work Location', 'Preferred work location (text, max 50 chars)'),
        ('13', 'Notice Period', 'Notice period in days (numeric, 0-90)'),
        ('14', 'Total Experience', 'Total years of experience (numeric, decimal allowed)'),
        ('15', 'Relevant Experience', 'Years of relevant experience (numeric, decimal allowed)'),
        ('16', 'Skills', 'Technical skills (comma-separated, max 500 chars)'),
        ('17', 'Status', 'Candidate status (dropdown: Submitted to Admin, With Admin, With Client, Selected, Onboarded, Rejected by Admin, Rejected by Client, Interview Scheduled, On Hold)'),
        ('18', 'Interview Time', 'Time of interview (time picker, HH:MM format)'),
        ('19', 'Remarks', 'Multi-line notes (textarea, max 1000 chars)'),
        ('20', 'Resume', 'Upload resume (file upload, PDF/DOCX only, max 5MB)'),
        ('21', 'Request ID', 'Linked resource request ID (auto-populated from parent request)')
    ]
    
    for num, name, desc in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = name
        row_cells[2].text = desc
    
    set_column_width(table, 0, 10)
    set_column_width(table, 1, 25)
    set_column_width(table, 2, 65)
    
    add_heading(doc, '5.2 Field Validation Rules', 2)
    doc.add_paragraph('All fields marked as "required" must be validated before form submission. Dropdown fields must enforce selection from predefined values only. File uploads must validate file type and size constraints.')
    
    # ============== 6. SCOPE & BOUNDARIES ==============
    add_heading(doc, '6. Scope & Boundaries', 1)
    
    add_heading(doc, '6.1 Phase 1: Core Platform (2 Weeks - Feb 16 to March 2, 2026)', 2)
    p = doc.add_paragraph()
    p.add_run('Included in Phase 1:').bold = True
    for item in [
        'User authentication and RBAC (2 roles)',
        'Dashboard with 6 core metrics and 2 charts',
        'Job profile CRUD',
        'Resource request workflow',
        'Recruiter pipeline with 21-field candidate form',
        'Admin review and client submission',
        'Onboarding workflow',
        'Exit management with backfill automation',
        'SOW tracker (manual entry)',
        'Basic reporting (export to Excel)',
        'Responsive UI for desktop browsers'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading(doc, '6.2 Phase 2: Advanced Features (Future - Post March 2026)', 2)
    p = doc.add_paragraph()
    p.add_run('Deferred to Phase 2:').bold = True
    for item in [
        'Client portal (direct login for clients to view profiles)',
        'Advanced analytics (predictive attrition modeling)',
        'Email integration (Gmail/Outlook sync)',
        'Calendar integration for interview scheduling',
        'Bulk upload (Excel import for candidates)',
        'Advanced search with filters (Boolean search)',
        'Mobile app (iOS/Android)',
        'Notifications (email/SMS alerts)',
        'Document version control for resumes',
        'API integrations (Jira, Workday)',
        'Multi-language support',
        'Custom report builder'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading(doc, '6.3 Out of Scope', 2)
    p = doc.add_paragraph()
    p.add_run('Not planned for any phase:').bold = True
    for item in [
        'Payroll management',
        'Time tracking or attendance',
        'Performance appraisals',
        'Learning management system',
        'Video interviewing platform'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # ============== 7. ASSUMPTIONS & DEPENDENCIES ==============
    add_heading(doc, '7. Assumptions & Dependencies', 1)
    
    add_heading(doc, '7.1 Technical Assumptions', 2)
    for item in [
        'Users have modern browsers (Chrome, Edge, Firefox - latest 2 versions)',
        'Stable internet connection (minimum 2 Mbps)',
        'Database hosting infrastructure is available',
        'Cloud storage for resume uploads is configured',
        'Development team has access to required tools and licenses'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading(doc, '7.2 Business Assumptions', 2)
    for item in [
        'Back Office team (5-7 users) is available for UAT',
        'HR leadership will approve the PRD within 3 business days',
        'Training will be conducted before go-live',
        'Legacy Excel data migration is out of scope (fresh start)',
        'SOW details will be entered manually by Admin'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading(doc, '7.3 External Dependencies', 2)
    for item in [
        'Database infrastructure provisioning (1-2 days)',
        'SSL certificate for HTTPS (provided by IT)',
        'Email server configuration for notifications (Phase 2)',
        'Client approvals are not required for Phase 1'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # ============== 8. SUCCESS CRITERIA & KPIs ==============
    add_heading(doc, '8. Success Criteria & KPIs', 1)
    
    add_heading(doc, '8.1 Quantitative Metrics', 2)
    doc.add_paragraph('The RMS will be considered successful if it achieves the following within 3 months of deployment:')
    
    for kpi in [
        '90% reduction in data entry errors (baseline: 30 errors/month → target: 3 errors/month)',
        '50% faster candidate-to-client submission time (baseline: 2 days → target: 1 day)',
        '100% audit trail coverage (all requests have logged source and communication)',
        'Zero manual SOW-to-headcount reconciliation effort (automated dashboard reporting)',
        '95% user adoption rate by Back Office team within 2 weeks of go-live',
        'System uptime of 99.5% (max 3.5 hours downtime/month)'
    ]:
        doc.add_paragraph(kpi, style='List Bullet')
    
    add_heading(doc, '8.2 Qualitative Metrics', 2)
    for metric in [
        'User satisfaction score ≥ 4/5 in post-deployment survey',
        'Positive feedback from client on profile quality (admin validation working)',
        'Reduced escalations to management for missing resume/profile data',
        'Improved confidence in headcount reporting for finance team'
    ]:
        doc.add_paragraph(metric, style='List Bullet')
    
    add_heading(doc, '8.3 User Adoption Targets', 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    data = [
        ('Milestone', 'Timeline', 'Target'),
        ('Training completion', 'Week 1 post-deployment', '100% of users'),
        ('Daily active usage', 'Week 2 post-deployment', '80% of users'),
        ('Full adoption (no Excel usage)', 'Week 4 post-deployment', '100% of users')
    ]
    
    for i, (col1, col2, col3) in enumerate(data):
        row_cells = table.rows[i].cells
        row_cells[0].text = col1
        row_cells[1].text = col2
        row_cells[2].text = col3
        
        if i == 0:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('--- END OF DOCUMENT ---', style='Heading 3')
    
    # Save
    output_path = 'C:/Users/parth/.gemini/antigravity/brain/39b25fbd-59cc-4ae8-8351-8dd232f82e33/final_prd.docx'
    doc.save(output_path)
    print(f'✅ final_prd.docx created successfully!')
    print(f'   Location: {output_path}')
    print(f'   - Complete PRD with all 8 sections')
    print(f'   - Text-based Table of Contents (manually convertible to table)')
    print(f'   - Proper table widths (10%, 25%, 65%)')
    print(f'   - Phase 1 and Phase 2 scope clearly defined')
    print(f'   - No approval sign-off section')

if __name__ == '__main__':
    create_complete_prd()
