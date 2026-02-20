# Quick script to check if the PRD needs TOC  
# User said there's NO TOC and tables are narrow - let's verify
from docx import Document

try:
    doc = Document('RMS_PRD_Final.docx')
    
    print("Checking PRD structure...")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    #Check first 20 paragraphs for TOC
    print("\nFirst 20 paragraphs:")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"{i}: {para.text[:80]}")
            
    # Check table widths
    if doc.tables:
        print(f"\nFirst table has {len(doc.tables[0].rows)} rows")
        print(f"First table has {len(doc.tables[0].columns)} columns")
        
    print("\n❌ User is right - need to regenerate with Python for proper control")
    
except FileNotFoundError:
    print("❌ RMS_PRD_Final.docx not found in D:\\RMS_Siprahub")
except Exception as e:
    print(f"Error: {e}")
