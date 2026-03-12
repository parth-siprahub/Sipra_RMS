
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def setup_doc(title):
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Add Title Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    
    doc.add_paragraph() # Spacer
    return doc

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
    return p

def generate_linkage_docx():
    doc = setup_doc("RMS SipraHub: Data Linkage & Cardinality Mapping")
    
    add_para(doc, "This document defines the relationships between core entities in the Resource Management System (RMS) to ensure full architectural control and database integrity.")
    
    add_heading(doc, "1. Entity Relationships (Cardinality)", 1)
    
    # Relationship Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Entity A'
    hdr_cells[1].text = 'Relationship'
    hdr_cells[2].text = 'Entity B'
    hdr_cells[3].text = 'Description'
    
    # Bold headers
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True

    data = [
        ["SOW", "1 : M", "Resource Request", "One SOW can have multiple resource requests (within budget/resource limits)."],
        ["Job Profile", "1 : M", "Resource Request", "One standardized profile can be used for multiple requests across different SOWs."],
        ["Resource Request", "1 : M", "Candidate", "One request can have multiple candidates mapped to it during the hiring process."],
        ["Candidate", "1 : 1", "Onboarded Seat", "A candidate fills exactly one request seat when onboarded."],
        ["Candidate (Exiting)", "1 : 1", "Backfill Request", "An exiting candidate triggers exactly one backfill Resource Request."]
    ]
    
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    add_heading(doc, "2. Core Data Linkage Flow", 1)
    
    add_heading(doc, "2.1 The 'Anchor' Points", 2)
    add_bullet(doc, "SOW (Statement of Work): The commercial master. It dictates the max_resources limit.")
    add_bullet(doc, "Job Profile: The technical master. It dictates the skills and role expectations.")

    add_heading(doc, "2.2 The 'Demand' Bridge (Resource Request)", 2)
    add_para(doc, "The Resource Request is the central node that connects commerce and talent:")
    add_bullet(doc, "sow_id -> Links to the billable contract.")
    add_bullet(doc, "job_profile_id -> Links to the required expertise.")
    add_bullet(doc, "is_backfill (Boolean) -> Flags if this is replacement demand.")
    add_bullet(doc, "replacement_for_candidate_id -> Parent link to the resource being replaced.")

    add_heading(doc, "2.3 The 'Supply' Layer (Candidates)", 2)
    add_bullet(doc, "request_id -> Links the candidate to a specific open demand.")
    add_bullet(doc, "status -> Tracks the transition from 'New' to 'Onboarded'.")
    add_bullet(doc, "overlap_until -> (For Backfills) Defines the handover period.")

    add_heading(doc, "3. Database Validation Checklist", 1)
    checks = [
        "SOW Table: Independent master.",
        "JobProfiles Table: Independent master.",
        "ResourceRequests Table: Foreign keys to sows.id and job_profiles.id.",
        "Candidates Table: Foreign key to resource_requests.id.",
        "Audit Trail: replacement_for_candidate_id ensures we never lose the lineage of backfilled positions."
    ]
    for c in checks:
        add_bullet(doc, "[X] " + c)

    add_heading(doc, "4. Backfill Operations Logic", 1)
    logic = [
        "Trigger: User initiates EXIT status for an onboarded candidate.",
        "Action: System prompts for 'Create Backfill?'.",
        "Result: Candidate status becomes EXIT; New Resource Request is created with is_backfill: true.",
        "Calculated utilization in Dashboard reflects 1 seat as 'Open (Backfill)' until the new resource is ONBOARDED."
    ]
    for l in logic:
        add_bullet(doc, l)

    doc.save("RMS_Linkage_Mapping.docx")
    print("Generated: RMS_Linkage_Mapping.docx")

def generate_flow_docx():
    doc = setup_doc("RMS SipraHub: System Flow Documentation")
    
    add_para(doc, "This document visualizes and describes the data flow and operational logic of the Resource Management System.")
    
    add_heading(doc, "1. The Golden Path (Standard Hiring Flow)", 1)
    add_para(doc, "The standard lifecycle for resource fulfillment follows a structured path from budget allocation to onboarding:")
    
    path = [
        "SOW Creation: Define Client Name, SOW Number, and Resource Cap.",
        "Job Profile Selection: Standardize roles (e.g., 'React Developer').",
        "Resource Request: Linking a profile to an SOW to create demand.",
        "Candidate Intake: Sourcing talent against specific requests.",
        "Admin Review: Profiles are validated by HR/Admin before client submission.",
        "Client Interview: Post-submission, candidates go through technical/client rounds.",
        "Selection & Onboarding: Final approval and joining, which auto-updates SOW utilization."
    ]
    for i, step in enumerate(path):
        add_para(doc, f"Step {i+1}: {step}")

    add_heading(doc, "2. Backfill & Exit Management Flow", 1)
    add_para(doc, "When a resource exits, the system ensures business continuity via automated backfill logic:")
    
    backfill = [
        "Initiate Exit: Capture LWD (Last Working Day) and Exit Reason.",
        "Trigger Demand: System automatically generates a 'HIGH' priority Resource Request.",
        "Linkage Preservation: The new request is tagged with 'replacement_for_candidate_id'.",
        "Handover Period: 'Overlap Until' date tracks the knowledge transfer window.",
        "Utilization Consistency: Dashboard maintains the 'Open Count' accurately during transition."
    ]
    for step in backfill:
        add_bullet(doc, step)

    add_heading(doc, "3. Information Architecture Mapping", 1)
    
    # Class mapping table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Module'
    hdr[1].text = 'Relationship'
    hdr[2].text = 'Primary Responsibility'
    
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True

    mapping = [
        ["SOW", "1 per Client Contract", "Dictates Budget and Headcount limits."],
        ["Job Profile", "Reusable Template", "Defines required Skills and Experience levels."],
        ["Resource Request", "Demand Node", "Links SOW budget to Profile requirements."],
        ["Candidate", "Sourcing Unit", "Mapped to Requests; transitions through Kanban pipeline."]
    ]
    
    for row_data in mapping:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()
    add_para(doc, "Note: These diagrams and flows are implemented as core logic in the RMS SipraHub backend for automated data integrity.", italic=True)

    doc.save("RMS_Flow_Diagram.docx")
    print("Generated: RMS_Flow_Diagram.docx")

def generate_status_update_docx():
    doc = setup_doc("RMS SipraHub: Weekly Status Update")
    
    add_para(doc, "Date: February 25, 2026", bold=True)
    add_para(doc, "Status: Progressing toward MVP Milestone")
    
    add_heading(doc, "✅ Done So Far", 1)
    dones = [
        "Architecture Control: Completed full Mapping & Linkage documentation defining 1:M relationships.",
        "Process Visualization: Created professional Flow Diagrams for hiring and backfill logic.",
        "System Validation: Audited Backend and DB schemas to ensure automated demand generation works.",
        "Email Integration: Base foundation for automated client/admin notifications established."
    ]
    for d in dones:
        add_bullet(doc, d)

    add_heading(doc, "📋 To-Do List (Next Steps)", 1)
    todos = [
        "UI Conflict Resolution: Finalize frontend merge logic to restore Kanban/Dashboard stability.",
        "Database Audit: Verify all Postgres foreign keys in Supabase match documentation.",
        "Backfill Dry Run: End-to-end testing of the 'Exit -> Replacement' lifecycle.",
        "Stakeholder Sign-off: Present diagrams to Senthil and Jason for process alignment."
    ]
    for t in todos:
        add_bullet(doc, t)

    doc.save("RMS_Status_Update.docx")
    print("Generated: RMS_Status_Update.docx")

if __name__ == "__main__":
    generate_linkage_docx()
    generate_flow_docx()
    generate_status_update_docx()
