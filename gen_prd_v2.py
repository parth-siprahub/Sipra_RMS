"""
PRD v2 Generator — Restructured per manager feedback
10 sections, no user stories, includes Tech/UI/DB sections
Meeting notes incorporated as professional content
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x3A, 0x6E, 0x9E)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
    return p

def set_col_widths(table, widths_pct):
    for row in table.rows:
        for i, pct in enumerate(widths_pct):
            row.cells[i].width = Inches(6.5 * pct / 100)

def create_prd_v2():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    
    # ============== TITLE PAGE ==============
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Product Requirements Document')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    run.font.name = 'Calibri'
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Resource Management System (RMS)')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    run.font.name = 'Calibri'
    
    for _ in range(2):
        doc.add_paragraph()
    
    meta_items = [
        ('Document Version:', '2.0'),
        ('Date:', 'February 19, 2026'),
        ('Project Timeline:', '2 Weeks (Feb 16 - March 2, 2026)'),
        ('Confidentiality:', 'Internal Use Only'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + ' ')
        r1.font.name = 'Calibri'
        r1.font.size = Pt(13)
        r1.bold = True
        r2 = p.add_run(value)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(13)
    
    doc.add_page_break()
    
    # ============== TABLE OF CONTENTS ==============
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc.add_run('TABLE OF CONTENTS')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    run.font.name = 'Calibri'
    run.bold = True
    
    doc.add_paragraph()
    
    toc_entries = [
        '1.  Executive Summary',
        '    1.1 Project Overview',
        '    1.2 Problem Statement',
        '    1.3 Business Objectives',
        '',
        '2.  User Personas',
        '    2.1 Back Office Recruiter',
        '    2.2 Admin (Profile Review & Client Submission)',
        '    2.3 Management / Leadership (Dashboard Viewer)',
        '    2.4 Client Portal User (Phase 2)',
        '',
        '3.  Functional Requirements',
        '    3.1 Authentication & RBAC',
        '    3.2 Dashboard & Real-Time Metrics',
        '    3.3 Job Profile Management',
        '    3.4 Resource Request Workflow',
        '    3.5 Recruiter Pipeline (21-Field Form)',
        '    3.6 Admin Review & Client Submission',
        '    3.7 Onboarding Workflow',
        '    3.8 Exit Management & Backfill',
        '    3.9 SOW Tracker',
        '',
        '4.  Candidate Fields Specification',
        '    4.1 21-Field Breakdown',
        '    4.2 Field Validation Rules',
        '',
        '5.  Technology & Architecture',
        '    5.1 Technology Stack',
        '    5.2 System Architecture',
        '    5.3 API Design',
        '    5.4 Security Model',
        '    5.5 Deployment & Infrastructure',
        '',
        '6.  UI Design Principles',
        '    6.1 Design System',
        '    6.2 Typography & Color Standards',
        '    6.3 Page Inventory',
        '    6.4 Responsive Strategy',
        '    6.5 Accessibility Requirements',
        '',
        '7.  Database Overview',
        '    7.1 Table Inventory',
        '    7.2 Key Relationships',
        '    7.3 Data Security',
        '',
        '8.  Scope & Boundaries',
        '    8.1 Phase 1: Core Platform (2 Weeks)',
        '    8.2 Phase 2: Advanced Features (Future)',
        '    8.3 Out of Scope',
        '',
        '9.  Assumptions & Dependencies',
        '    9.1 Technical Assumptions',
        '    9.2 Business Assumptions',
        '    9.3 External Dependencies',
        '',
        '10. Success Criteria & KPIs',
        '    10.1 Quantitative Metrics',
        '    10.2 Qualitative Metrics',
        '    10.3 User Adoption Targets',
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        if entry and not entry.startswith('    '):
            run.bold = True
    
    doc.add_page_break()
    
    # ============== 1. EXECUTIVE SUMMARY ==============
    add_heading(doc, '1. Executive Summary', 1)
    
    add_heading(doc, '1.1 Project Overview', 2)
    add_para(doc, 'The Resource Management System (RMS) is a comprehensive web-based platform designed to replace the current manual Excel-based workflow for managing staff augmentation operations. The system will centralize tracking of resource requests, candidate sourcing, admin review, client submission, onboarding, and lifecycle management from initial request through exit.')
    
    add_heading(doc, '1.2 Problem Statement', 2)
    add_para(doc, 'The organization currently manages approximately 300+ active resources using multiple Excel spreadsheets (Resource Data, SOW Tracker, Category Data, Recruiter Pipeline). This manual process results in:')
    problems = [
        'Data inconsistencies due to manual entry (e.g., role name variations like "Java Developer" vs "Java Dev")',
        'No centralized profile storage — requires email searches to retrieve candidate resumes',
        'Client visibility into high attrition rates when sharing Excel files directly',
        'Difficult billing reconciliation with client systems due to Jira username mapping issues',
        'No audit trail for request origins (email vs chat) or client communications',
        'Inability to scale beyond 300-500 resources without exponential manual effort',
        'Backfill requests created manually with no linkage to original exits or rejections',
    ]
    for p in problems:
        add_bullet(doc, p)
    
    add_heading(doc, '1.3 Business Objectives', 2)
    objectives = [
        'Centralize all staff augmentation operations in one secure, web-based platform',
        'Reduce data entry errors by 90% through structured forms and validation rules',
        'Provide real-time visibility into resource pipeline, onboarded headcount, and attrition trends',
        'Streamline recruiter-to-admin handoff for client profile submissions',
        'Automate backfill request creation for exits and client rejections',
        'Enable scalability to support 500+ concurrent resources without performance degradation',
        'Create comprehensive audit trail for all request sources and client communications',
    ]
    for o in objectives:
        add_bullet(doc, o)
    
    # ============== 2. USER PERSONAS ==============
    doc.add_page_break()
    add_heading(doc, '2. User Personas', 1)
    
    add_heading(doc, '2.1 Back Office Recruiter', 2)
    add_para(doc, 'Role:', bold=True)
    add_para(doc, 'Handles candidate sourcing, screening, and initial submission to admin.')
    add_para(doc, 'Responsibilities:', bold=True)
    recruiter_resp = [
        'Create resource requests when client requirements are received',
        'Source candidates from vendors (WRS, GFM) or internal database',
        'Add candidates with complete 21-field profile form',
        'Upload candidate resumes (PDF/DOCX)',
        'Update candidate status through pipeline (Interview Scheduled, Submitted to Admin, etc.)',
        'Mark requests as "With Admin" when ready for profile review',
    ]
    for r in recruiter_resp:
        add_bullet(doc, r)
    add_para(doc, 'Pain Points (Current):', bold=True)
    pains = [
        'Manually copying data across multiple Excel sheets',
        'No dropdown validations — frequent typos in role names and statuses',
        'Cannot track which recruiter owns which candidate',
        'Searching emails to find candidate resumes submitted weeks ago',
    ]
    for p in pains:
        add_bullet(doc, p)
    
    add_heading(doc, '2.2 Admin (Profile Review & Client Submission)', 2)
    add_para(doc, 'Role:', bold=True)
    add_para(doc, 'Reviews recruiter-submitted profiles and sends them to clients.')
    add_para(doc, 'Responsibilities:', bold=True)
    admin_resp = [
        'View all requests marked "With Admin"',
        'Validate candidate profiles (email format, JD alignment, resume quality)',
        'Reject profiles back to recruiter with feedback if needed',
        'Generate professional email templates for client submission',
        'Log communication details (date sent, client contact)',
        'Update status to "With Client" after submission',
        'Process client selections and rejections',
        'Track L1/L2 interview response statuses from clients',
        'Manage profile rejection flow and trigger backfill when needed',
    ]
    for r in admin_resp:
        add_bullet(doc, r)
    
    add_heading(doc, '2.3 Management / Leadership (Dashboard Viewer)', 2)
    add_para(doc, 'Role:', bold=True)
    add_para(doc, 'Reviews real-time status of the resource pipeline and operational health.')
    add_para(doc, 'Responsibilities:', bold=True)
    mgmt_resp = [
        'View dashboard with aggregated metrics and charts',
        'Monitor resource status across all stages (sourcing, with admin, with client, onboarded)',
        'Track attrition trends and onboarding velocity',
        'Access role-wise and technology-wise breakdowns for planning',
    ]
    for r in mgmt_resp:
        add_bullet(doc, r)
    
    add_heading(doc, '2.4 Client Portal User (Phase 2)', 2)
    add_para(doc, 'Future persona — not part of Phase 1 scope. Will enable clients to directly log in, view submitted profiles, and provide feedback. See Section 8.2 for Phase 2 features.', italic=True)
    
    # ============== 3. FUNCTIONAL REQUIREMENTS ==============
    doc.add_page_break()
    add_heading(doc, '3. Functional Requirements', 1)
    
    add_heading(doc, '3.1 Authentication & Role-Based Access Control', 2)
    auth_reqs = [
        'User login with email and password',
        'Three roles: Back Office (Recruiter), Admin, and Management (read-only)',
        'Permission-based feature access (e.g., only Admin can "Send to Client")',
        'Session persistence across page refresh',
        'Logout functionality',
    ]
    for r in auth_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.2 Dashboard & Real-Time Metrics', 2)
    dash_reqs = [
        'Total requests metric',
        'Onboarded count',
        'Awaiting onboarding count (status = "With Client")',
        'To be shared count (status = "With Admin")',
        'Role-wise breakdown (bar/pie chart)',
        'Technology distribution chart',
        'Status filter dropdown',
        'Attrition rate trend graph (Phase 1 basic, Phase 2 advanced — more dashboards planned)',
        'Average time to onboard metric',
        'Date range filter',
    ]
    for r in dash_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.3 Job Profile Management', 2)
    jp_reqs = [
        'Create job profile (role name, technology, experience)',
        'Edit job profile',
        'Delete job profile (with validation — cannot delete if linked to requests)',
        'List job profiles (paginated)',
        'Duplicate validation (same role name cannot exist twice)',
    ]
    for r in jp_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.4 Resource Request Workflow', 2)
    req_reqs = [
        'Create request with auto-generated ID (REQ-YYYYMMDD-XXX)',
        'Select job profile from dropdown',
        'Capture request source (Email/Chat)',
        'Set priority (Urgent/High/Medium/Low)',
        'Support multi-position requests',
        'View request list (paginated)',
        'Filter requests by status, priority, job profile',
        'Search by request ID or candidate name',
    ]
    for r in req_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.5 Recruiter Pipeline (21-Field Form)', 2)
    add_para(doc, 'Note: Field list to be confirmed with team during Sprint 1.', italic=True)
    pipe_reqs = [
        'Add candidate with 21 fields (see Section 4)',
        'Owner assignment (recruiter dropdown)',
        'Vendor field (WRS/GFM/Internal)',
        'Interview date & time picker',
        'Candidate status dropdown (9 options: Submitted to Admin, With Admin, With Client, Selected, Onboarded, Rejected by Admin, Rejected by Client, Interview Scheduled, On Hold)',
        'CTC fields (current/expected)',
        'Location fields (current/work location)',
        'Notice period field',
        'Remarks (multi-line text)',
        'Upload resume (PDF/DOCX, max 5MB)',
        'Edit candidate details (all 21 fields)',
        'View candidate list for request',
        'Update request status to "With Admin"',
    ]
    for r in pipe_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.6 Admin Review & Client Submission', 2)
    admin_reqs = [
        'View requests "With Admin"',
        'Validation checklist (email format, JD alignment, resume quality)',
        'Download candidate profile as PDF',
        'Reject to Back Office with reason',
        'Generate email template for client',
        'Update status to "With Client"',
        'Log communication details (date, client contact, type)',
        'Track L1/L2 interview response status from client',
        'Handle profile rejection flow (client rejects → option for replacement)',
    ]
    for r in admin_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.7 Onboarding Workflow', 2)
    onboard_reqs = [
        'Mark as onboarded',
        'Capture billing start date',
        'Capture client email ID',
        'Capture client Jira username',
        'Update status to "Onboarded"',
        'Client rejection flow (with reason)',
        'Replacement required choice (backfill yes/no)',
        'Auto-create backfill on rejection if replacement needed',
    ]
    for r in onboard_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.8 Exit Management & Backfill', 2)
    exit_reqs = [
        'Process exit (capture reason & last working day)',
        'Exit reason dropdown (6 options: Better Offer, Personal, Performance, Client End, Project End, Other)',
        'Last working day date picker',
        'Exit notes (multi-line text)',
        'Replacement required choice',
        'Auto-create backfill on exit if replacement needed',
        'Update status to "Exit"',
    ]
    for r in exit_reqs:
        add_bullet(doc, r)
    
    add_heading(doc, '3.9 SOW Tracker', 2)
    sow_reqs = [
        'Manual SOW entry form',
        'Link SOW to request IDs',
        'View SOW list (table — default filter: active SOWs)',
        'Display request count per SOW (active SOW instances = number of resources)',
        'Filter SOW by status (Active/Expired/All)',
    ]
    for r in sow_reqs:
        add_bullet(doc, r)
    
    # ============== 4. CANDIDATE FIELDS ==============
    doc.add_page_break()
    add_heading(doc, '4. Candidate Fields Specification', 1)
    
    add_heading(doc, '4.1 21-Field Breakdown', 2)
    add_para(doc, 'Complete field specification for candidate form:')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Field Name'
    hdr[2].text = 'Description / Validation Rules'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    fields = [
        ('1', 'Owner', 'Recruiter assigned (dropdown from users table)'),
        ('2', 'Date', 'Date candidate added (auto-populated, DD-MMM-YYYY)'),
        ('3', 'Vendor', 'Sourcing vendor (dropdown: WRS, GFM, Internal)'),
        ('4', 'Interview Date', 'Scheduled interview date (date picker)'),
        ('5', 'Candidate Name', 'Full name (text, max 100 chars, required)'),
        ('6', 'Email', 'Email address (email validation, required)'),
        ('7', 'Phone', 'Contact number (10 digits, numeric)'),
        ('8', 'Current Company', 'Current employer name (text, max 100 chars)'),
        ('9', 'Current CTC', 'Current annual salary (numeric, in lakhs)'),
        ('10', 'Expected CTC', 'Expected annual salary (numeric, in lakhs)'),
        ('11', 'Current Location', 'Current city/location (text, max 50 chars)'),
        ('12', 'Work Location', 'Preferred work location (text, max 50 chars)'),
        ('13', 'Notice Period', 'Notice period in days (numeric, 0-90)'),
        ('14', 'Total Experience', 'Total years of experience (numeric, decimal allowed)'),
        ('15', 'Relevant Experience', 'Years of relevant experience (numeric, decimal allowed)'),
        ('16', 'Skills', 'Technical skills (comma-separated, max 500 chars)'),
        ('17', 'Status', 'Candidate status (dropdown: 9 options — see Section 3.5)'),
        ('18', 'Interview Time', 'Time of interview (time picker, HH:MM format)'),
        ('19', 'Remarks', 'Multi-line notes (textarea, max 1000 chars)'),
        ('20', 'Resume', 'Upload resume (PDF/DOCX only, max 5MB)'),
        ('21', 'Request ID', 'Linked resource request ID (auto-populated from parent)'),
    ]
    for num, name, desc in fields:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = name
        row[2].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    set_col_widths(table, [10, 25, 65])
    
    add_heading(doc, '4.2 Field Validation Rules', 2)
    add_para(doc, 'All fields marked as "required" must be validated before form submission. Dropdown fields must enforce selection from predefined values only. File uploads must validate file type and size constraints.')
    
    # ============== 5. TECHNOLOGY & ARCHITECTURE ==============
    doc.add_page_break()
    add_heading(doc, '5. Technology & Architecture', 1)
    
    add_heading(doc, '5.1 Technology Stack', 2)
    
    stack_table = doc.add_table(rows=1, cols=3)
    stack_table.style = 'Table Grid'
    hdr = stack_table.rows[0].cells
    hdr[0].text = 'Layer'
    hdr[1].text = 'Technology'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    stack = [
        ('Frontend', 'React 18 + TypeScript + Vite', 'Single Page Application with type safety'),
        ('Backend', 'Python FastAPI + Uvicorn', 'Async REST API server'),
        ('Database', 'PostgreSQL (Supabase)', 'Cloud-hosted relational database'),
        ('Auth', 'Supabase Auth / JWT', 'Authentication and session management'),
        ('Storage', 'Supabase Storage', 'Resume file uploads (PDF/DOCX)'),
        ('Deployment', 'TBD (Vercel / Railway / VPS)', 'Hosting and CI/CD'),
    ]
    for layer, tech, purpose in stack:
        row = stack_table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = purpose
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    set_col_widths(stack_table, [15, 35, 50])
    
    add_heading(doc, '5.2 System Architecture', 2)
    add_para(doc, 'The system follows a standard 3-tier architecture:')
    add_bullet(doc, 'Presentation Layer: React SPA communicating via REST APIs')
    add_bullet(doc, 'Application Layer: FastAPI backend handling business logic, validation, and authorization')
    add_bullet(doc, 'Data Layer: Supabase PostgreSQL with Row Level Security (RLS)')
    add_para(doc, 'All communication between frontend and backend is over HTTPS. The backend exposes versioned REST endpoints (/api/v1/...). File uploads are stored in Supabase Storage with signed URLs for secure access.')
    
    add_heading(doc, '5.3 API Design', 2)
    add_bullet(doc, 'RESTful conventions with proper HTTP methods (GET, POST, PUT, DELETE)')
    add_bullet(doc, 'Versioned endpoints: /api/v1/')
    add_bullet(doc, 'Pydantic models for request/response validation')
    add_bullet(doc, 'Standardized error responses with error codes')
    add_bullet(doc, 'Pagination support for all list endpoints')
    
    add_heading(doc, '5.4 Security Model', 2)
    add_bullet(doc, 'JWT-based authentication with token refresh')
    add_bullet(doc, 'Role-based access control (RBAC) enforced at API level')
    add_bullet(doc, 'Row Level Security (RLS) policies on Supabase tables')
    add_bullet(doc, 'CORS configuration restricted to frontend origin')
    add_bullet(doc, 'Input validation on all API endpoints')
    add_bullet(doc, 'File upload restrictions (type, size)')
    
    add_heading(doc, '5.5 Deployment & Infrastructure', 2)
    add_bullet(doc, 'Supabase project for database, auth, and storage')
    add_bullet(doc, 'Environment-based configuration (.env files)')
    add_bullet(doc, 'CI/CD pipeline (TBD — GitHub Actions or similar)')
    add_bullet(doc, 'SSL/TLS for all connections')
    add_bullet(doc, 'Health check endpoints for monitoring')
    
    # ============== 6. UI DESIGN PRINCIPLES ==============
    doc.add_page_break()
    add_heading(doc, '6. UI Design Principles', 1)
    
    add_heading(doc, '6.1 Design System', 2)
    add_para(doc, 'The RMS interface must maintain visual consistency across all screens. Key design principles:')
    add_bullet(doc, 'Consistent component library — all forms, tables, buttons, and modals use shared components')
    add_bullet(doc, 'Unified theming — single theme applied across the entire application')
    add_bullet(doc, 'Design tokens for spacing, border-radius, and shadows')
    add_bullet(doc, 'Reusable layout templates for list views, detail views, and forms')
    
    add_heading(doc, '6.2 Typography & Color Standards', 2)
    add_bullet(doc, 'Primary font family to be selected (e.g., Inter, Roboto, or equivalent modern sans-serif)')
    add_bullet(doc, 'Consistent font sizing hierarchy: headings, body, captions, labels')
    add_bullet(doc, 'Primary brand color palette with semantic colors (success, warning, error, info)')
    add_bullet(doc, 'Sufficient color contrast for readability (WCAG AA compliance)')
    add_para(doc, 'Note: Final color palette and font selection to be confirmed during UI discussion with team.', italic=True)
    
    add_heading(doc, '6.3 Page Inventory', 2)
    add_para(doc, '8 core screens in Phase 1:')
    
    screens_table = doc.add_table(rows=1, cols=3)
    screens_table.style = 'Table Grid'
    hdr = screens_table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Screen'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    screens = [
        ('1', 'Login', 'Email + password authentication'),
        ('2', 'Dashboard', 'Metrics, charts, status overview'),
        ('3', 'Job Profiles', 'CRUD for job profile management'),
        ('4', 'Resource Requests', 'Request list, create, filter, search'),
        ('5', 'Recruiter Pipeline', 'Candidate 21-field form, list, status updates'),
        ('6', 'Admin Review', 'Profile validation, email generation, client submission'),
        ('7', 'Lifecycle (Onboard/Exit)', 'Onboarding capture, exit processing, backfill'),
        ('8', 'SOW Tracker', 'SOW entry, linking, table view'),
    ]
    for num, screen, purpose in screens:
        row = screens_table.add_row().cells
        row[0].text = num
        row[1].text = screen
        row[2].text = purpose
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    set_col_widths(screens_table, [10, 25, 65])
    
    add_heading(doc, '6.4 Responsive Strategy', 2)
    add_bullet(doc, 'Desktop-first design (primary users are on laptops/desktops)')
    add_bullet(doc, 'Minimum supported viewport: 1280px width')
    add_bullet(doc, 'Tablet support (1024px) for dashboard viewing')
    add_bullet(doc, 'Mobile support deferred to Phase 2')
    
    add_heading(doc, '6.5 Accessibility Requirements', 2)
    add_bullet(doc, 'Keyboard navigation for all interactive elements')
    add_bullet(doc, 'ARIA labels for screen reader support')
    add_bullet(doc, 'Form field labels and error messages clearly visible')
    add_bullet(doc, 'Minimum touch target size of 44x44px for clickable elements')
    
    # ============== 7. DATABASE OVERVIEW ==============
    doc.add_page_break()
    add_heading(doc, '7. Database Overview', 1)
    
    add_heading(doc, '7.1 Table Inventory', 2)
    add_para(doc, 'The RMS database consists of 11 core tables:')
    
    db_table = doc.add_table(rows=1, cols=3)
    db_table.style = 'Table Grid'
    hdr = db_table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Table Name'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    tables = [
        ('1', 'users', 'Application users with role assignments (Back Office, Admin, Management)'),
        ('2', 'job_profiles', 'Job role definitions (role name, technology, experience range)'),
        ('3', 'resource_requests', 'Client resource requests with auto-generated IDs, priority, source'),
        ('4', 'candidates', '21-field candidate profiles linked to requests'),
        ('5', 'candidate_status_history', 'Audit trail of candidate status changes with timestamps'),
        ('6', 'resumes', 'Resume file metadata and storage references'),
        ('7', 'admin_reviews', 'Admin validation checklist results and actions'),
        ('8', 'communications', 'Client communication logs (date, contact, type, notes)'),
        ('9', 'onboarding', 'Onboarding details (billing date, client email, Jira username)'),
        ('10', 'exits', 'Exit records (reason, LWD, notes, replacement flag)'),
        ('11', 'sow', 'Statement of Work entries linked to request IDs'),
    ]
    for num, name, purpose in tables:
        row = db_table.add_row().cells
        row[0].text = num
        row[1].text = name
        row[2].text = purpose
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    set_col_widths(db_table, [10, 25, 65])
    
    add_heading(doc, '7.2 Key Relationships', 2)
    add_bullet(doc, 'resource_requests → job_profiles (Many-to-One: each request references one job profile)')
    add_bullet(doc, 'candidates → resource_requests (Many-to-One: multiple candidates per request)')
    add_bullet(doc, 'candidates → users (Many-to-One: owner/recruiter assignment)')
    add_bullet(doc, 'candidate_status_history → candidates (One-to-Many: status change audit trail)')
    add_bullet(doc, 'resumes → candidates (One-to-One: one active resume per candidate)')
    add_bullet(doc, 'admin_reviews → candidates (One-to-Many: review history)')
    add_bullet(doc, 'communications → resource_requests (One-to-Many: communication log per request)')
    add_bullet(doc, 'onboarding → candidates (One-to-One: onboarding record per candidate)')
    add_bullet(doc, 'exits → candidates (One-to-One: exit record per candidate)')
    add_bullet(doc, 'sow → resource_requests (Many-to-Many: SOW linked to multiple requests)')
    
    add_heading(doc, '7.3 Data Security', 2)
    add_bullet(doc, 'Row Level Security (RLS) policies on all tables')
    add_bullet(doc, 'Role-based read/write access at database level')
    add_bullet(doc, 'Sensitive fields (CTC, phone) restricted to authorized roles')
    add_bullet(doc, 'All timestamps stored in UTC (TIMESTAMPTZ)')
    add_bullet(doc, 'Soft delete pattern for audit compliance (records marked as deleted, not removed)')
    
    # ============== 8. SCOPE & BOUNDARIES ==============
    doc.add_page_break()
    add_heading(doc, '8. Scope & Boundaries', 1)
    
    add_heading(doc, '8.1 Phase 1: Core Platform (2 Weeks — Feb 16 to March 2, 2026)', 2)
    add_para(doc, 'Included in Phase 1:', bold=True)
    phase1 = [
        'User authentication and RBAC (3 roles: Back Office, Admin, Management)',
        'Dashboard with 6 core metrics and 2 charts',
        'Job profile CRUD',
        'Resource request workflow',
        'Recruiter pipeline with 21-field candidate form',
        'Admin review and client submission',
        'Onboarding workflow',
        'Exit management with backfill automation',
        'SOW tracker (manual entry with active filter default)',
        'Basic reporting (export to Excel)',
        'Responsive UI for desktop browsers',
    ]
    for p in phase1:
        add_bullet(doc, p)
    
    add_heading(doc, '8.2 Phase 2: Advanced Features (Future — Post March 2026)', 2)
    add_para(doc, 'Deferred to Phase 2:', bold=True)
    phase2 = [
        'Client portal (direct login for clients to view profiles)',
        'Advanced analytics and dashboards (predictive attrition modeling, custom report builder)',
        'Email integration (Gmail/Outlook sync)',
        'Calendar integration for interview scheduling',
        'Bulk upload (Excel import for candidates)',
        'Advanced search with filters (Boolean search)',
        'Mobile app (iOS/Android)',
        'Notifications (email/SMS alerts)',
        'Document version control for resumes',
        'API integrations (Jira, Workday)',
        'Multi-language support',
    ]
    for p in phase2:
        add_bullet(doc, p)
    
    add_heading(doc, '8.3 Out of Scope', 2)
    oos = [
        'Payroll management',
        'Time tracking or attendance',
        'Performance appraisals',
        'Learning management system',
        'Video interviewing platform',
    ]
    for o in oos:
        add_bullet(doc, o)
    
    # ============== 9. ASSUMPTIONS ==============
    doc.add_page_break()
    add_heading(doc, '9. Assumptions & Dependencies', 1)
    
    add_heading(doc, '9.1 Technical Assumptions', 2)
    tech_assum = [
        'Users have modern browsers (Chrome, Edge, Firefox — latest 2 versions)',
        'Stable internet connection (minimum 2 Mbps)',
        'Database hosting infrastructure is available (Supabase)',
        'Cloud storage for resume uploads is configured',
        'Development team has access to required tools and licenses',
    ]
    for a in tech_assum:
        add_bullet(doc, a)
    
    add_heading(doc, '9.2 Business Assumptions', 2)
    biz_assum = [
        'Back Office team (5-7 users) is available for UAT',
        'HR leadership will approve the PRD within 3 business days',
        'Training will be conducted before go-live',
        'Legacy Excel data migration is out of scope (fresh start)',
        'SOW details will be entered manually by Admin',
    ]
    for a in biz_assum:
        add_bullet(doc, a)
    
    add_heading(doc, '9.3 External Dependencies', 2)
    ext_deps = [
        'Database infrastructure provisioning (1-2 days)',
        'SSL certificate for HTTPS (provided by IT)',
        'Email server configuration for notifications (Phase 2)',
        'Client approvals are not required for Phase 1',
    ]
    for d in ext_deps:
        add_bullet(doc, d)
    
    # ============== 10. SUCCESS CRITERIA ==============
    doc.add_page_break()
    add_heading(doc, '10. Success Criteria & KPIs', 1)
    
    add_heading(doc, '10.1 Quantitative Metrics', 2)
    add_para(doc, 'The RMS will be considered successful if it achieves the following within 3 months of deployment:')
    quant = [
        '90% reduction in data entry errors (baseline: 30 errors/month → target: 3 errors/month)',
        '50% faster candidate-to-client submission time (baseline: 2 days → target: 1 day)',
        '100% audit trail coverage (all requests have logged source and communication)',
        'Zero manual SOW-to-headcount reconciliation effort (automated dashboard reporting)',
        '95% user adoption rate by Back Office team within 2 weeks of go-live',
        'System uptime of 99.5% (max 3.5 hours downtime/month)',
    ]
    for q in quant:
        add_bullet(doc, q)
    
    add_heading(doc, '10.2 Qualitative Metrics', 2)
    qual = [
        'User satisfaction score >= 4/5 in post-deployment survey',
        'Positive feedback from client on profile quality (admin validation working)',
        'Reduced escalations to management for missing resume/profile data',
        'Improved confidence in headcount reporting for finance team',
    ]
    for q in qual:
        add_bullet(doc, q)
    
    add_heading(doc, '10.3 User Adoption Targets', 2)
    
    adopt_table = doc.add_table(rows=1, cols=3)
    adopt_table.style = 'Table Grid'
    hdr = adopt_table.rows[0].cells
    hdr[0].text = 'Milestone'
    hdr[1].text = 'Timeline'
    hdr[2].text = 'Target'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    adoption = [
        ('Training completion', 'Week 1 post-deployment', '100% of users'),
        ('Daily active usage', 'Week 2 post-deployment', '80% of users'),
        ('Full adoption (no Excel)', 'Week 4 post-deployment', '100% of users'),
    ]
    for milestone, timeline, target in adoption:
        row = adopt_table.add_row().cells
        row[0].text = milestone
        row[1].text = timeline
        row[2].text = target
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    set_col_widths(adopt_table, [35, 35, 30])
    
    # ============== END ==============
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run('— END OF DOCUMENT —')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    output = r'C:/Users/parth/.gemini/antigravity/brain/39b25fbd-59cc-4ae8-8351-8dd232f82e33/final_prd_v2.docx'
    doc.save(output)
    print(f'PRD v2 generated: {output}')
    print('  10 sections, TOC, meeting notes incorporated, no user stories')

if __name__ == '__main__':
    create_prd_v2()
