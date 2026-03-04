"""
RMS PRD v2.0 — DOCX Generator
Produces a professionally formatted Word document at D:\RMS_Siprahub\docs\RMS_PRD_v2_Updated.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colour palette ──
NAVY = RGBColor(0x1F, 0x47, 0x88)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT_GREEN = RGBColor(0x1A, 0x7A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = '1F4788'
ALT_ROW_BG = 'F2F5FB'
TAG_BG = 'E8F4FD'

OUT_PATH = os.path.join(os.path.dirname(__file__), 'RMS_PRD_v2_Updated.docx')


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg=HEADER_BG):
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
    style_header_row(tbl.rows[0])

    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK_GREY
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    # Widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    return tbl


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(13)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
        run.font.size = Pt(11)
    return p


def para(doc, text, bold=False, italic=False, size=10, color=DARK_GREY, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = space_after
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    return p


def tag_para(doc, text, tag_text='v2.0 UPDATE'):
    p = doc.add_paragraph()
    tag_run = p.add_run(f'[{tag_text}] ')
    tag_run.font.bold = True
    tag_run.font.size = Pt(9)
    tag_run.font.color.rgb = RGBColor(0x0C, 0x54, 0x60)
    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = DARK_GREY
    return p


def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_GREY

    # ================================================================
    # COVER / TITLE BLOCK
    # ================================================================
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Product Requirements Document')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Resource Management System (RMS)')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    
    doc.add_paragraph()

    meta_items = [
        ('Document Version:', '2.0 (Updated)'),
        ('Original Version:', '1.0 — February 16, 2026 — Jaicind Santhibhavan'),
        ('Update Date:', 'March 2, 2026'),
        ('Update Author:', 'Parth P (RMS Engineering Lead)'),
        ('Project Timeline:', '2 Weeks (Feb 16 – March 2, 2026)'),
        ('Confidentiality:', 'Internal Use Only'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = p.add_run(label + ' ')
        lbl.font.bold = True
        lbl.font.size = Pt(10)
        lbl.font.color.rgb = MID_GREY
        val = p.add_run(value)
        val.font.size = Pt(10)
        val.font.color.rgb = DARK_GREY

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run('This document is an additive update to the original PRD. All changes are tagged [v2.0 UPDATE] or [v2.0 NEW].')
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MID_GREY

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    h1(doc, 'Table of Contents')
    toc_items = [
        '1. Executive Summary',
        '    1.1 Project Overview',
        '    1.2 Problem Statement',
        '    1.3 Business Objectives',
        '2. User Personas',
        '    2.1 Back Office Recruiter',
        '    2.2 Admin (Profile Review & Client Submission)',
        '    2.3 Management Persona (v2.0 NEW)',
        '    2.4 Client Portal User (Phase 2)',
        '3. User Stories & Acceptance Criteria (US-001 to US-010)',
        '4. Functional Requirements',
        '    4.1 Authentication & Role-Based Access Control',
        '    4.2 Dashboard & Real-Time Metrics',
        '    4.3 Job Profile Management',
        '    4.4 Resource Request Workflow',
        '    4.5 Recruiter Pipeline (21-Field Form)',
        '    4.6 Admin Review & Client Submission',
        '    4.7 Onboarding Workflow',
        '    4.8 Exit Management & Backfill',
        '    4.9 SOW Tracker (Manual Entry)',
        '    4.10 Vendor Management (v2.0 NEW)',
        '5. Candidate Fields Specification',
        '6. Scope & Boundaries',
        '    6A. Phase 1 — Delivered Features Summary (v2.0 NEW)',
        '7. Assumptions & Dependencies',
        '8. Success Criteria & KPIs',
        '9. Database Schema — Actual (v2.0 NEW)',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            if 'v2.0 NEW' in item:
                run.font.color.rgb = RGBColor(0x0C, 0x54, 0x60)
                run.font.bold = True
            else:
                run.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    h1(doc, '1. Executive Summary')

    h2(doc, '1.1 Project Overview')
    para(doc, 'The Resource Management System (RMS) is a comprehensive web-based platform designed to replace the current manual Excel-based workflow for managing staff augmentation operations. The system will centralize tracking of resource requests, candidate sourcing, admin review, client submission, onboarding, and lifecycle management from initial request through exit.')

    h2(doc, '1.2 Problem Statement')
    para(doc, 'The organization currently manages approximately 300+ active resources using multiple Excel spreadsheets (Resource Data, SOW Tracker, Category Data, Recruiter Pipeline). This manual process results in:')
    problems = [
        'Data inconsistencies due to manual entry (e.g., role name variations like "Java Developer" vs "Java Dev")',
        'No centralized profile storage — requires email searches to retrieve candidate resumes',
        'Client visibility into high attrition rates when sharing Excel files directly',
        'Difficult billing reconciliation with client systems due to Jira username mapping issues',
        'No audit trail for request origins (email vs chat) or client communications',
        'Inability to scale beyond 300–500 resources without exponential manual effort',
        'Backfill requests created manually with no linkage to original exits/rejections',
    ]
    for p_text in problems:
        bullet(doc, p_text)

    h2(doc, '1.3 Business Objectives')
    para(doc, 'The RMS aims to:')
    objectives = [
        'Centralize all staff augmentation operations in one secure, web-based platform',
        'Reduce data entry errors by 90% through structured forms and validation rules',
        'Provide real-time visibility into resource pipeline, onboarded headcount, and attrition trends',
        'Streamline recruiter-to-admin handoff for client profile submissions',
        'Automate backfill request creation for exits and client rejections',
        'Enable scalability to support 500+ concurrent resources without performance degradation',
        'Create comprehensive audit trail for all request sources and client communications',
    ]
    for o in objectives:
        bullet(doc, o)

    # ================================================================
    # 2. USER PERSONAS
    # ================================================================
    doc.add_page_break()
    h1(doc, '2. User Personas')

    h2(doc, '2.1 Back Office Recruiter')
    para(doc, 'Role: Handles candidate sourcing, screening, and initial submission to admin', bold=True)
    para(doc, 'Responsibilities:')
    for r in [
        'Create resource requests when client requirements are received',
        'Source candidates from vendors (WRS, GFM) or internal database',
        'Add candidates with complete 21-field profile form',
        'Upload candidate resumes (PDF/DOCX)',
        'Update candidate status through pipeline',
        'Mark requests as "With Admin" when ready for profile review',
    ]:
        bullet(doc, r)
    para(doc, 'Pain Points (Current):', bold=True)
    for pp in [
        'Manually copying data across multiple Excel sheets',
        'No dropdown validations — frequent typos in role names and statuses',
        'Cannot track which recruiter owns which candidate',
        'Searching emails to find candidate resumes submitted weeks ago',
    ]:
        bullet(doc, pp)

    h2(doc, '2.2 Admin (Profile Review & Client Submission)')
    para(doc, 'Role: Reviews recruiter-submitted profiles and sends them to clients', bold=True)
    para(doc, 'Responsibilities:')
    for r in [
        'View all requests marked "With Admin"',
        'Validate candidate profiles (email format, JD alignment, resume quality)',
        'Reject profiles back to recruiter with feedback if needed',
        'Generate professional email templates for client submission',
        'Log communication details (date sent, client contact)',
        'Update status to "With Client" after submission',
        'Process client selections and rejections',
    ]:
        bullet(doc, r)

    h2(doc, '2.3 Management Persona')
    tag_para(doc, 'New persona added based on stakeholder calls with Raja PV and Senthil Natarajan.', 'v2.0 NEW')
    para(doc, 'Role: Senior leadership visibility into pipeline and delivery health', bold=True)
    para(doc, 'Responsibilities:')
    for r in [
        'View dashboards showing overall pipeline health',
        'Review onboarded headcount by role and technology',
        'Monitor attrition trends and SOW alignment',
        'Review vendor performance metrics',
        'Access read-only view of resource statuses across all requests',
    ]:
        bullet(doc, r)
    para(doc, 'Confirmed during demo on February 23, 2026. MANAGEMENT role is implemented in the system with read-only dashboard access.', italic=True, size=9, color=MID_GREY)

    h2(doc, '2.4 Client Portal User (Phase 2)')
    para(doc, 'Future persona — not part of Phase 1 scope. See Section 6.2 for Phase 2 features.')
    para(doc, 'Set of people (client-side) who need to review and see status — should be able to see dashboards and view resource in which stage.', italic=True, size=9, color=MID_GREY)

    # ================================================================
    # 3. USER STORIES
    # ================================================================
    doc.add_page_break()
    h1(doc, '3. User Stories & Acceptance Criteria')

    user_stories = [
        ('US-001: Resource Request Creation',
         'As a Back Office Recruiter, I want to create resource requests with auto-generated IDs so that I can track client requirements systematically.',
         [
             'Request ID is auto-generated (format: REQ-YYYYMMDD-XXX)',
             'Job profile is selected from dropdown',
             'Request source (Email/Chat) is captured',
             'Priority (Urgent/High/Medium/Low) can be set',
             'Multi-position requests are supported',
             'Request is saved and visible in list view',
         ], []),
        ('US-002: Candidate Addition (21-Field Form)',
         'As a Back Office Recruiter, I want to add candidates with complete profiles so that Admin can review all necessary information.',
         [
             'All 21 fields can be filled (see Section 5)',
             'Resume can be uploaded (PDF/DOCX, max 5MB)',
             'Owner (recruiter) is assigned from dropdown',
             'Candidate status defaults to "New"',
             'Candidate is linked to request ID',
             'Form validates required fields before save',
         ], []),
        ('US-003: Admin Profile Review',
         'As an Admin, I want to review recruiter-submitted profiles so that I only send quality candidates to clients.',
         [
             'Can view all requests "With Admin"',
             'Can see complete candidate profile with resume',
             'Can reject to recruiter with reason',
             'Can approve and proceed to client submission',
             'Validation checklist is visible (email, JD match, resume quality)',
         ],
         ['L1/L2 interview feedback statuses are captured as part of the candidate record.']),
        ('US-004: Client Submission',
         'As an Admin, I want to generate professional emails for client submission so that the process is standardized.',
         [
             'Email template is auto-generated with candidate details',
             'Can download candidate profile as PDF',
             'Can log communication date and client contact',
             'Status updates to "With Client" after submission',
             'Audit trail is maintained',
         ], []),
        ('US-005: Onboarding Workflow',
         'As an Admin, I want to capture onboarding details so that billing and client systems are synchronized.',
         [
             'Can mark candidate as onboarded',
             'Billing start date is captured',
             'Client email ID is captured',
             'Client Jira username is captured',
             'Status updates to "Onboarded"',
             'Dashboard metrics update in real-time',
         ], []),
        ('US-006: Exit Processing',
         'As an Admin, I want to process exits with reasons so that attrition trends can be analyzed.',
         [
             'Can select exit reason from dropdown (6 options)',
             'Last working day is captured',
             'Exit notes can be added',
             'Replacement required flag can be set',
             'Backfill request is auto-created if replacement needed',
             'Status updates to "Exit"',
         ], []),
        ('US-007: Dashboard Visibility',
         'As a user, I want to see real-time metrics so that I can track pipeline health.',
         [
             'Total requests count is displayed',
             'Onboarded count is displayed',
             'Awaiting onboarding count (status = "With Client") is shown',
             'To be shared count (status = "With Admin") is shown',
             'Role-wise breakdown chart is displayed',
             'Technology distribution is visible',
         ], []),
        ('US-008: SOW Tracker',
         'As an Admin, I want to track SOW details so that headcount can be reconciled with contracts.',
         [
             'Can manually enter SOW details',
             'Can link request IDs to SOW',
             'Can view SOW list in table format (default: Active SOWs only)',
             'Can see request count per SOW',
         ],
         [
             'SOW is mandatory-linked to each Resource Request (confirmed: Raja PV)',
             'Active instances on an SOW represent onboarded resources under that contract',
         ]),
        ('US-009: Backfill Automation',
         'As an Admin, I want backfill requests to be auto-created so that replacements are not forgotten.',
         [
             'Backfill is created when exit has "replacement required" = Yes',
             'Backfill is created when client rejection has "replacement required" = Yes',
             'Backfill request inherits job profile from original request',
             'Backfill is linked to original request ID',
         ],
         [
             'For backfill scenarios, the same SOW carries forward from the original request',
             'Overlap period (1–2 weeks) is trackable via the overlap_until date field',
         ]),
        ('US-010: Communication Logging',
         'As an Admin, I want to log all client communications so that there is an audit trail.',
         [
             'Can log communication date',
             'Can log client contact person',
             'Can log communication type (email/call/meeting)',
             'All logs are timestamped and linked to request',
         ], []),
    ]

    for us_title, us_desc, ac_items, v2_items in user_stories:
        h3(doc, us_title)
        para(doc, us_desc, italic=True)
        para(doc, 'Acceptance Criteria:', bold=True)
        for ac in ac_items:
            bullet(doc, ac)
        for v2 in v2_items:
            tag_para(doc, v2)

    # ================================================================
    # 4. FUNCTIONAL REQUIREMENTS
    # ================================================================
    doc.add_page_break()
    h1(doc, '4. Functional Requirements')

    # 4.1
    h2(doc, '4.1 Authentication & Role-Based Access Control')
    for b in [
        'User login with email and password',
        'Permission-based feature access (e.g., only Admin can "Send to Client"; Management has read-only)',
        'Session persistence across page refresh',
        'Logout functionality',
    ]:
        bullet(doc, b)
    tag_para(doc, 'Three roles: Back Office (Recruiter), Admin, and Management')

    # 4.2
    h2(doc, '4.2 Dashboard & Real-Time Metrics')
    for b in [
        'Total requests metric',
        'Onboarded count',
        'Awaiting onboarding count (status = "With Client")',
        'To be shared count (status = "With Admin")',
        'Role-wise breakdown (bar/pie chart)',
        'Technology distribution chart',
        'Status filter dropdown',
        'Attrition rate trend graph (Phase 1 basic, Phase 2 advanced)',
        'Average time to onboard metric',
        'Date range filter',
    ]:
        bullet(doc, b)
    tag_para(doc, 'Vendor performance breakdown (sourcing by vendor)')
    tag_para(doc, 'SOW utilization summary')
    tag_para(doc, 'Granular candidate pipeline status breakdown (per HR-approved statuses)')

    # 4.3
    h2(doc, '4.3 Job Profile Management')
    for b in [
        'Create job profile (role name, technology, experience)',
        'Edit job profile',
        'Delete job profile (with validation — cannot delete if linked to requests)',
        'List job profiles (paginated)',
        'Duplicate validation (same role name cannot exist twice)',
    ]:
        bullet(doc, b)

    # 4.4
    h2(doc, '4.4 Resource Request Workflow')
    for b in [
        'Create request with auto-generated ID (REQ-YYYYMMDD-XXX)',
        'Select job profile from dropdown',
        'Capture request source (Email/Chat)',
        'Set priority (Urgent/High/Medium/Low)',
        'Support multi-position requests',
        'View request list (paginated)',
        'Filter requests by status, priority, job profile',
        'Search by request ID or candidate name',
    ]:
        bullet(doc, b)
    tag_para(doc, 'SOW linkage is mandatory — each resource request must be associated with an SOW')
    tag_para(doc, 'Request type field: New / Backfill')
    tag_para(doc, 'overlap_until date field to track overlap period during backfill scenarios')

    # 4.5
    h2(doc, '4.5 Recruiter Pipeline (21-Field Form)')
    for b in [
        'Add candidate with 21 fields (see Section 5)',
        'Owner assignment (recruiter dropdown)',
        'Vendor field (dropdown — see Section 4.10 for vendor master)',
        'Interview date & time picker',
        'CTC fields (current/expected)',
        'Location fields (current/work location)',
        'Notice period field',
        'Remarks (multi-line text)',
        'Upload resume (PDF/DOCX, max 5MB)',
        'Edit candidate details (all 21 fields)',
        'View candidate list for request',
        'Update request status to "With Admin"',
    ]:
        bullet(doc, b)
    tag_para(doc, 'Candidate status dropdown with HR-approved statuses: NEW, SCREENING_DONE, L1_COMPLETED, L1_REJECTED, L2_COMPLETED, L2_REJECTED, SELECTED, WITH_CLIENT, ONBOARDED, EXIT')
    tag_para(doc, 'L1/L2 interview feedback fields capture round-specific assessment')
    tag_para(doc, 'Kanban view showing candidates grouped by pipeline status', 'v2.0 NEW')

    # 4.6
    h2(doc, '4.6 Admin Review & Client Submission')
    for b in [
        'View requests "With Admin"',
        'Validation checklist (email format, JD alignment, resume quality)',
        'Download candidate profile as PDF',
        'Reject to Back Office with reason',
        'Generate email template for client',
        'Update status to "With Client"',
        'Log communication details (date, client contact, type)',
    ]:
        bullet(doc, b)

    # 4.7
    h2(doc, '4.7 Onboarding Workflow')
    for b in [
        'Mark as onboarded',
        'Capture billing start date',
        'Capture client email ID',
        'Capture client Jira username',
        'Update status to "Onboarded"',
        'Client rejection flow (with reason)',
        'Replacement required choice',
        'Auto-create backfill on rejection if replacement needed',
    ]:
        bullet(doc, b)

    # 4.8
    h2(doc, '4.8 Exit Management & Backfill')
    for b in [
        'Process exit (capture reason & last working day)',
        'Exit reason dropdown (6 options: Better Offer, Personal, Performance, Client End, Project End, Other)',
        'Last working day date picker',
        'Exit notes (multi-line text)',
        'Replacement required choice',
        'Auto-create backfill on exit if replacement needed',
        'Update status to "Exit"',
    ]:
        bullet(doc, b)
    tag_para(doc, 'When backfill is triggered: same SOW carries forward; overlap_until date is captured to track transition period')

    # 4.9
    h2(doc, '4.9 SOW Tracker (Manual Entry)')
    for b in [
        'Manual SOW entry form',
        'Link SOW to request IDs',
        'View SOW list (table)',
        'Display request count per SOW',
    ]:
        bullet(doc, b)
    tag_para(doc, 'Default filter: Active SOWs only (per Jaicind\'s instruction)')
    tag_para(doc, 'Active instances count on each SOW = number of currently Onboarded resources')
    tag_para(doc, 'SOW is a mandatory field on Resource Requests (confirmed: Raja PV)')

    # 4.10
    h2(doc, '4.10 Vendor Management')
    tag_para(doc, 'Vendor is a master module — not merely a dropdown field.', 'v2.0 NEW')
    for b in [
        'CRUD operations for vendor master (Admin only)',
        'Vendors included: WRS, GFM, Internal (SipraHub direct sourcing), and additional Anton vendors',
        'Vendor is linked to each candidate record',
        'Vendor field drives dashboard analytics on sourcing distribution',
        'Admin can add/edit/deactivate vendors without a code change',
    ]:
        bullet(doc, b)

    # ================================================================
    # 5. CANDIDATE FIELDS
    # ================================================================
    doc.add_page_break()
    h1(doc, '5. Candidate Fields Specification')
    h2(doc, '5.1 21-Field Breakdown')

    field_rows = [
        ('1', 'Owner', 'Recruiter assigned (dropdown from users table)'),
        ('2', 'Date', 'Date candidate added (auto-populated, DD-MMM-YYYY)'),
        ('3', 'Vendor', 'Sourcing vendor (dropdown from Vendor master)'),
        ('4', 'Interview Date', 'Scheduled interview date (date picker)'),
        ('5', 'Candidate Name', 'Full name (text, max 100 chars, required)'),
        ('6', 'Email', 'Email address (email validation, required)'),
        ('7', 'Phone', 'Contact number (10 digits, numeric)'),
        ('8', 'Current Company', 'Current employer name (text, max 100 chars)'),
        ('9', 'Current CTC', 'Current annual salary (numeric, in lakhs)'),
        ('10', 'Expected CTC', 'Expected annual salary (numeric, in lakhs)'),
        ('11', 'Current Location', 'Current city/location (text, max 50 chars)'),
        ('12', 'Work Location', 'Preferred work location (text, max 50 chars)'),
        ('13', 'Notice Period', 'Notice period in days (numeric, 0–90)'),
        ('14', 'Total Experience', 'Total years of experience (numeric, decimal)'),
        ('15', 'Relevant Experience', 'Years of relevant experience (numeric, decimal)'),
        ('16', 'Skills', 'Technical skills (comma-separated, max 500 chars)'),
        ('17', 'Status', '[v2.0] NEW, SCREENING_DONE, L1_COMPLETED, L1_REJECTED, L2_COMPLETED, L2_REJECTED, SELECTED, WITH_CLIENT, ONBOARDED, EXIT'),
        ('18', 'Interview Time', 'Time of interview (time picker, HH:MM)'),
        ('19', 'Remarks', 'Multi-line notes (textarea, max 1000 chars)'),
        ('20', 'Resume', 'Upload resume (PDF/DOCX only, max 5MB)'),
        ('21', 'Request ID', 'Linked resource request ID (auto-populated)'),
    ]
    add_table(doc, ['#', 'Field Name', 'Description / Validation Rules'], field_rows, [1.5, 4, 12])

    h2(doc, '5.2 Field Validation Rules')
    para(doc, 'All fields marked as "required" must be validated before form submission. Dropdown fields must enforce selection from predefined values only. File uploads must validate file type and size constraints.')
    tag_para(doc, 'The candidate status field must enforce the 10 HR-approved values listed above. Legacy statuses from v1.0 are superseded and replaced.')

    # ================================================================
    # 6. SCOPE & BOUNDARIES
    # ================================================================
    doc.add_page_break()
    h1(doc, '6. Scope & Boundaries')

    h2(doc, '6.1 Phase 1: Core Platform — ✅ COMPLETE')
    tag_para(doc, 'Phase 1 delivered on schedule. Demo to management on February 23, 2026.', 'v2.0 UPDATE')
    para(doc, 'Included in Phase 1 (all delivered):')
    for b in [
        'User authentication and RBAC (3 roles: Admin, Recruiter, Management)',
        'Dashboard with core metrics and bar/pie charts',
        'Job profile CRUD',
        'Resource request workflow with REQ-YYYYMMDD-XXX auto-IDs',
        'Recruiter pipeline with 21-field candidate form',
        'Kanban view of candidate pipeline',
        'Admin review and client submission',
        'Onboarding workflow (billing date, client Jira ID)',
        'Exit management with backfill automation',
        'SOW tracker (manual entry, linked to requests)',
        'Communication logs (per request and per candidate)',
        'Vendor Management module (CRUD master)',
        'L1/L2 interview feedback capture',
        'Resume upload to Supabase Storage',
    ]:
        bullet(doc, b)

    h2(doc, '6.2 Phase 2: Advanced Features (Post March 2026)')
    tag_para(doc, 'Phase 2 scope expanded with 8 new items.', 'v2.0 UPDATE')
    para(doc, 'Deferred to Phase 2:')
    phase2_items = [
        'Client portal (direct login for clients)',
        'Advanced analytics (predictive attrition modeling)',
        'Email integration (Gmail/Outlook sync)',
        'Calendar integration for interview scheduling',
        'Bulk upload (Excel import for candidates)',
        'Advanced search with filters (Boolean search)',
        'Mobile app (iOS/Android)',
        'Notifications (email/SMS alerts)',
        'Document version control for resumes',
        'API integrations (Jira, Workday, LinkedIn/Job Boards)',
        'Multi-language support',
        'Custom report builder',
    ]
    for b in phase2_items:
        bullet(doc, b)

    para(doc, 'New Phase 2 items (v2.0):', bold=True)
    new_p2 = [
        'Payroll Management — Jira time dump CSV import + payroll calculation engine',
        'Exit Management Dedicated Page — full exit workflow with analytics',
        'Onboarding Workflow Page — dedicated onboarding UX',
        'Production Hosting & Deployment (Railway/Fly.io)',
        'PR Review Workflow Setup — GitHub branch protection + manager approval',
        'User Training Materials — SOPs and guides',
        'Resume Parsing (AI-powered extraction)',
        'LinkedIn / Job Board API Integration',
        'Advanced Dashboard Analytics — deeper breakdowns, trends, PDF export',
    ]
    for b in new_p2:
        tag_para(doc, b, 'v2.0 NEW')

    h2(doc, '6.3 Out of Scope')
    for b in [
        'Time tracking or attendance',
        'Performance appraisals',
        'Learning management system',
        'Video interviewing platform',
    ]:
        bullet(doc, b)
    tag_para(doc, 'Payroll management reclassified from Out-of-Scope to Phase 2 (per Raja PV and HR team direction).', 'v2.0 UPDATE')

    # ================================================================
    # 6A. PHASE 1 DELIVERED FEATURES SUMMARY
    # ================================================================
    doc.add_page_break()
    h1(doc, '6A. Phase 1 — Delivered Features Summary')
    tag_para(doc, 'New section added in v2.0 to document what was actually built and delivered.', 'v2.0 NEW')

    h2(doc, 'Overview')
    para(doc, 'Phase 1 was successfully delivered and demo\'d to management (Raja PV and Senthil Natarajan) on February 23, 2026, completing on schedule within the 2-week sprint.')

    h2(doc, 'Technology Stack')
    add_table(doc, ['Layer', 'Technology'], [
        ('Backend', 'Python FastAPI + Uvicorn (async)'),
        ('Frontend', 'React 18 + TypeScript + Vite'),
        ('Database', 'Supabase (PostgreSQL)'),
        ('Authentication', 'Supabase Auth (JWT, ES256)'),
        ('File Storage', 'Supabase Storage (resume uploads)'),
        ('API Count', '23 endpoints'),
    ], [4, 12])

    doc.add_paragraph()
    h2(doc, 'Database — 6 Tables')
    add_table(doc, ['Table', 'Purpose'], [
        ('profiles', 'User accounts + role (ADMIN, RECRUITER, MANAGEMENT)'),
        ('job_profiles', 'Job role master (role name, technology, experience)'),
        ('resource_requests', 'Resource requests with auto-generated REQ IDs'),
        ('candidates', '21-field candidate records linked to requests'),
        ('communication_logs', 'Audit trail — per request and per candidate'),
        ('sows', 'Statement of Work master, linked to requests'),
    ], [5, 12])

    doc.add_paragraph()
    h2(doc, 'API Endpoints (23)')
    api_groups = [
        '/auth/login, /auth/me',
        '/job-profiles/ (CRUD — admin write only)',
        '/requests/ (CRUD + status transitions)',
        '/candidates/ (CRUD, resume upload)',
        '/sows/ (CRUD — admin write only)',
        '/logs/ (create + list)',
        '/dashboard/metrics',
        '/health',
    ]
    for a in api_groups:
        bullet(doc, a)

    doc.add_paragraph()
    h2(doc, 'E2E Testing')
    para(doc, '17 E2E test flows verified via Playwright covering: Login (3 roles), Dashboard, Requests CRUD, Candidates CRUD, Kanban view, Resume upload, SOW CRUD, Communication logs, Admin status transitions, Vendor management.')

    doc.add_paragraph()
    h2(doc, 'Key Stakeholder Decisions (Locked)')
    decisions = [
        ('1', 'SOW mandatory-linked to Resource Requests', 'Raja PV'),
        ('2', 'L1/L2 interview feedback captured in candidate record', 'HR Team'),
        ('3', 'Backfill: same SOW carries forward; overlap_until tracks transition', 'Raja PV'),
        ('4', 'SOW default view: Active only', 'Jaicind'),
        ('5', 'Candidate statuses: 10 HR-approved values (NEW → EXIT)', 'HR Team'),
        ('6', 'Vendor field: master module with Anton + SipraHub sourcing', 'HR Team'),
        ('7', 'SOW entry: manual (no automation in Phase 1)', 'HR Team'),
        ('8', 'Payroll → Phase 2: Jira time dump CSV import + calculation', 'Raja PV'),
        ('9', 'PR review workflow: daily commits → manager approval → merge', 'Jaicind'),
    ]
    add_table(doc, ['#', 'Decision', 'Confirmed By'], decisions, [1.5, 12, 4])

    # ================================================================
    # 7. ASSUMPTIONS & DEPENDENCIES
    # ================================================================
    doc.add_page_break()
    h1(doc, '7. Assumptions & Dependencies')

    h2(doc, '7.1 Technical Assumptions')
    for b in [
        'Users have modern browsers (Chrome, Edge, Firefox — latest 2 versions)',
        'Stable internet connection (minimum 2 Mbps)',
        'Database hosting infrastructure is available',
        'Cloud storage for resume uploads is configured',
        'Development team has access to required tools and licenses',
    ]:
        bullet(doc, b)

    h2(doc, '7.2 Business Assumptions')
    for b in [
        'Back Office team (5–7 users) is available for UAT',
        'HR leadership will approve the PRD within 3 business days',
        'Training will be conducted before go-live',
        'Legacy Excel data migration is out of scope (fresh start)',
        'SOW details will be entered manually by Admin',
    ]:
        bullet(doc, b)

    h2(doc, '7.3 External Dependencies')
    for b in [
        'Database infrastructure provisioning (1–2 days)',
        'SSL certificate for HTTPS (provided by IT)',
        'Email server configuration for notifications (Phase 2)',
        'Client approvals are not required for Phase 1',
    ]:
        bullet(doc, b)

    # ================================================================
    # 8. SUCCESS CRITERIA & KPIs
    # ================================================================
    h1(doc, '8. Success Criteria & KPIs')

    h2(doc, '8.1 Quantitative Metrics')
    para(doc, 'The RMS will be considered successful if it achieves the following within 3 months of deployment:')
    for b in [
        '90% reduction in data entry errors (baseline: 30 errors/month → target: 3)',
        '50% faster candidate-to-client submission time (baseline: 2 days → target: 1 day)',
        '100% audit trail coverage (all requests have logged source and communication)',
        'Zero manual SOW-to-headcount reconciliation effort',
        '95% user adoption rate by Back Office team within 2 weeks of go-live',
        'System uptime of 99.5% (max 3.5 hours downtime/month)',
    ]:
        bullet(doc, b)

    h2(doc, '8.2 Qualitative Metrics')
    for b in [
        'User satisfaction score ≥ 4/5 in post-deployment survey',
        'Positive feedback from client on profile quality',
        'Reduced escalations to management for missing resume/profile data',
        'Improved confidence in headcount reporting for finance team',
    ]:
        bullet(doc, b)

    h2(doc, '8.3 User Adoption Targets')
    add_table(doc, ['Milestone', 'Timeline', 'Target'], [
        ('Training completion', 'Week 1 post-deployment', '100% of users'),
        ('Daily active usage', 'Week 2 post-deployment', '80% of users'),
        ('Full adoption (no Excel)', 'Week 4 post-deployment', '100% of users'),
    ], [6, 6, 5])

    # ================================================================
    # 9. DATABASE SCHEMA
    # ================================================================
    doc.add_page_break()
    h1(doc, '9. Database Schema (Actual — v2.0)')
    tag_para(doc, 'The original PRD projected 11 tables. The actual implementation uses 6 tables — scope was rationalized during development in consultation with the HR team.', 'v2.0 NEW')

    h2(doc, 'Actual Tables (6)')
    add_table(doc, ['Table', 'Key Columns', 'Notes'], [
        ('profiles', 'id, email, role, full_name', 'Supabase Auth-linked'),
        ('job_profiles', 'id, role_name, technology, experience_level', 'Unique constraint on role_name'),
        ('resource_requests', 'id, req_id, job_profile_id, sow_id, status, type, overlap_until, priority', 'sow_id is a mandatory FK'),
        ('candidates', 'id, request_id, candidate_name, email, status, vendor_id, resume_url, l1/l2 feedback +15 fields', 'Status constrained to 10 HR-approved values'),
        ('communication_logs', 'id, request_id, candidate_id, log_type, log_date, contact_person', 'Dual-link: per request OR per candidate'),
        ('sows', 'id, sow_number, client_name, start_date, end_date, is_active, headcount', 'is_active drives default filter'),
    ], [4, 8, 5])

    doc.add_paragraph()
    h2(doc, 'Phase 2 Database Additions (Planned)')
    for b in [
        'vendors — Vendor master (normalize from string field)',
        'payroll_records — Jira dump imports + calculated payroll',
        'exit_records — Dedicated exit management table',
        'onboarding_records — Dedicated onboarding workflow table',
        'user_training_logs — Training completion tracking',
    ]:
        bullet(doc, b)

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    para(doc, 'END OF DOCUMENT v2.0', bold=True, size=9, color=MID_GREY)
    para(doc, 'Original PRD v1.0 authored by Jaicind Santhibhavan, February 16, 2026.', italic=True, size=9, color=MID_GREY)
    para(doc, 'Updates authored by Parth P (RMS Engineering Lead), March 2, 2026.', italic=True, size=9, color=MID_GREY)

    # ── Save ──
    doc.save(OUT_PATH)
    print(f'✅ RMS_PRD_v2_Updated.docx saved to: {OUT_PATH}')
    print(f'   Pages: ~12 (estimated)')


if __name__ == '__main__':
    build_document()
